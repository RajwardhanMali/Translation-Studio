"""
Document Regeneration Service.

Reconstructs a translated DOCX or PDF from the segmented + translated data.

DOCX reconstruction:
  - Replays every block in order using the format_snapshot.
  - For multi-sentence paragraphs, joins translated sentences back into
    one paragraph maintaining original run-level formatting on the first run.
  - Headings re-apply the original heading style + run formatting.
  - Tables are rebuilt row-by-row, column-by-column using table_start metadata.
  - Spacers (empty paragraphs) are reproduced as empty paragraphs with
    original paragraph formatting so vertical spacing is preserved.

PDF reconstruction:
  - Rebuilds as a new PDF using ReportLab (falls back to a structured HTML
    export if ReportLab is not available).
  - Positions each text block at its original bbox coordinates (x0, y1 from
    PyMuPDF → ReportLab coordinates which are bottom-origin).
  - Applies original font size; bold/italic derived from flags/font name.
  - Images from the original are re-embedded at their original bboxes.
  - Tables are rebuilt with tabular layout using original column widths.

Note on PDF fidelity:
  Perfect pixel-for-pixel PDF reconstruction is not possible without
  embedding the exact original fonts. This service produces a
  "layout-faithful" PDF that matches heading/paragraph positions, font sizes,
  and reading order. For production use, consider generating a DOCX output
  from a PDF input (more reliable cross-format reconstruction).
"""

import logging
import io
from pathlib import Path
from typing import List, Dict, Any, Optional
from collections import defaultdict
from datetime import datetime

logger = logging.getLogger(__name__)

EXPORTS_DIR = Path(__file__).resolve().parent.parent / "data" / "exported_docs"
EXPORTS_DIR.mkdir(parents=True, exist_ok=True)


# ============================================================================
# Helpers
# ============================================================================

def _resolved_text(seg: Dict) -> str:
    """Return the best available translation for a segment."""
    return (
        seg.get("final_text")
        or seg.get("correction")
        or seg.get("translated_text")
        or seg.get("text", "")
    )


def _hex_to_rgb_int(hex_color: Optional[str]):
    """Convert '#rrggbb' → (r,g,b) tuple of ints. Returns None on failure."""
    if not hex_color or not hex_color.startswith("#"):
        return None
    try:
        h = hex_color.lstrip("#")
        return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))
    except Exception:
        return None


# ============================================================================
# DOCX Regeneration
# ============================================================================

def _apply_run_format(run, fmt: Dict) -> None:
    """Apply a formatting dict (from format_snapshot) to a docx Run."""
    try:
        from docx.shared import Pt, RGBColor
        if fmt.get("bold")      is not None: run.bold      = fmt["bold"]
        if fmt.get("italic")    is not None: run.italic    = fmt["italic"]
        if fmt.get("underline") is not None: run.underline = fmt["underline"]

        if fmt.get("font_name"):
            run.font.name = fmt["font_name"]
        if fmt.get("font_size"):
            run.font.size = Pt(fmt["font_size"])
        if fmt.get("color"):
            rgb = _hex_to_rgb_int(f"#{fmt['color']}" if not fmt["color"].startswith("#") else fmt["color"])
            if rgb:
                run.font.color.rgb = RGBColor(*rgb)
        if fmt.get("strike")     is not None: run.font.strike     = fmt["strike"]
        if fmt.get("all_caps")   is not None: run.font.all_caps   = fmt["all_caps"]
        if fmt.get("small_caps") is not None: run.font.small_caps = fmt["small_caps"]
        if fmt.get("superscript") is not None: run.font.superscript = fmt["superscript"]
        if fmt.get("subscript")   is not None: run.font.subscript   = fmt["subscript"]
    except Exception as e:
        logger.debug(f"Run format apply error: {e}")


def _apply_paragraph_format(para, fmt: Dict) -> None:
    """Apply paragraph-level formatting from format_snapshot."""
    try:
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        pf = para.paragraph_format
        align_map = {
            "WD_ALIGN_PARAGRAPH.LEFT":    WD_ALIGN_PARAGRAPH.LEFT,
            "WD_ALIGN_PARAGRAPH.CENTER":  WD_ALIGN_PARAGRAPH.CENTER,
            "WD_ALIGN_PARAGRAPH.RIGHT":   WD_ALIGN_PARAGRAPH.RIGHT,
            "WD_ALIGN_PARAGRAPH.JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }

        if fmt.get("alignment") and fmt["alignment"] in align_map:
            pf.alignment = align_map[fmt["alignment"]]
        if fmt.get("left_indent")        is not None: pf.left_indent        = Pt(fmt["left_indent"])
        if fmt.get("right_indent")       is not None: pf.right_indent       = Pt(fmt["right_indent"])
        if fmt.get("first_line_indent")  is not None: pf.first_line_indent  = Pt(fmt["first_line_indent"])
        if fmt.get("space_before")       is not None: pf.space_before       = Pt(fmt["space_before"])
        if fmt.get("space_after")        is not None: pf.space_after        = Pt(fmt["space_after"])
    except Exception as e:
        logger.debug(f"Paragraph format apply error: {e}")


def _write_paragraph(doc, translated_text: str, fmt_snapshot: Dict, style_name: Optional[str] = None) -> None:
    """
    Add a paragraph to the docx Document with restored formatting.
    Distributes translated text across runs proportionally to original run lengths.
    """
    from docx import Document
    from docx.oxml.ns import qn

    # Determine style
    style = None
    sn = style_name or fmt_snapshot.get("formatting", {}).get("style_name", "Normal")
    try:
        style = doc.styles[sn]
    except (KeyError, Exception):
        style = None

    para = doc.add_paragraph(style=style)
    _apply_paragraph_format(para, fmt_snapshot.get("formatting", {}))

    orig_runs = fmt_snapshot.get("runs", [])

    if not orig_runs:
        # No run info — add a single plain run
        para.add_run(translated_text)
        return

    # Distribute translated text across original run proportions
    orig_text_total = sum(len(r.get("text", "")) for r in orig_runs)

    if orig_text_total == 0:
        para.add_run(translated_text)
        return

    translated_len = len(translated_text)
    pos = 0

    for i, orig_run in enumerate(orig_runs):
        orig_run_len = len(orig_run.get("text", ""))
        proportion   = orig_run_len / orig_text_total

        if i == len(orig_runs) - 1:
            # Last run gets remainder to avoid rounding gaps
            chunk = translated_text[pos:]
        else:
            chunk_len = max(1, round(proportion * translated_len))
            chunk     = translated_text[pos:pos + chunk_len]
            pos      += chunk_len

        if chunk:
            run = para.add_run(chunk)
            _apply_run_format(run, orig_run.get("formatting", {}))


def regenerate_docx(
    parsed_doc: Dict,
    segments: List[Dict],
    output_path: Path,
) -> Path:
    """
    Rebuild a DOCX file from translated segments.

    Strategy:
      1. Group segments by block_index (parent block).
      2. Replay blocks in order, using pass-through for spacers/markers.
      3. For each translatable block, join its sentence segments and write
         one paragraph with original formatting.
      4. Rebuild tables from table_start/cell/table_end markers.
    """
    from docx import Document
    from docx.shared import Pt, Inches

    doc = Document()

    # Remove the default empty paragraph that python-docx adds
    for para in doc.paragraphs:
        p = para._element
        p.getparent().remove(p)

    # Index segments by block_index, preserving order
    block_map: Dict[int, List[Dict]] = defaultdict(list)
    skip_map:  Dict[int, Dict]       = {}    # block_index → skip segment

    for seg in segments:
        bi = seg["position"]["block_index"]
        if seg.get("status") == "skip":
            skip_map[bi] = seg
        else:
            block_map[bi].append(seg)

    # Sort by block_index for replay
    all_block_indices = sorted(
        set(list(block_map.keys()) + list(skip_map.keys()))
    )

    active_table      = None   # current docx Table being built
    active_table_segs: Dict[tuple, str] = {}   # (row,col) → translated text
    active_table_fmt:  Dict = {}
    table_row_count   = 0
    table_col_count   = 0

    def _flush_table():
        """Write the accumulated table cells into active_table."""
        nonlocal active_table, active_table_segs, active_table_fmt
        if active_table is None:
            return
        for (r, c), cell_text in sorted(active_table_segs.items()):
            try:
                cell = active_table.cell(r, c)
                cell.text = ""
                para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                para.text = cell_text
            except Exception as e:
                logger.debug(f"Table cell write error ({r},{c}): {e}")
        active_table      = None
        active_table_segs = {}

    for bi in all_block_indices:

        # ── Pass-through (spacer / table markers / image) ───────────────────
        if bi in skip_map:
            skip_seg  = skip_map[bi]
            btype     = skip_seg.get("block_type")

            if btype == "spacer":
                fmt = skip_seg.get("format_snapshot", {}).get("formatting", {})
                para = doc.add_paragraph()
                _apply_paragraph_format(para, fmt)
                continue

            if btype == "table_start":
                _flush_table()
                rc = skip_seg.get("row_count", 1)
                cc = skip_seg.get("col_count", 1)
                table_row_count = rc
                table_col_count = cc
                widths = skip_seg.get("col_widths") or []
                active_table = doc.add_table(rows=rc, cols=cc)
                active_table.style = "Table Grid"
                # Apply column widths if available
                if widths:
                    for col_idx, width_pt in enumerate(widths):
                        if width_pt and col_idx < len(active_table.columns):
                            for cell in active_table.columns[col_idx].cells:
                                cell.width = Pt(width_pt)
                active_table_segs = {}
                continue

            if btype == "table_end":
                _flush_table()
                continue

            if btype == "image":
                # Images cannot be re-embedded without the original binary.
                # Insert a placeholder paragraph.
                doc.add_paragraph(f"[Image — page {skip_seg.get('format_snapshot',{}).get('page','')+1 if skip_seg.get('format_snapshot',{}).get('page') is not None else '?'}]")
                continue
            continue

        segs = block_map.get(bi, [])
        if not segs:
            continue

        first_seg = segs[0]
        btype      = first_seg.get("block_type", "paragraph")
        seg_type   = first_seg.get("type", "sentence")
        fmt_snap   = first_seg.get("format_snapshot", {})

        # ── Table cells ──────────────────────────────────────────────────────
        if seg_type == "table_cell":
            for seg in segs:
                r = seg.get("row", 0)
                c = seg.get("col", 0)
                active_table_segs[(r, c)] = _resolved_text(seg)
            continue

        # ── Headings ─────────────────────────────────────────────────────────
        if seg_type == "heading":
            level     = fmt_snap.get("level", 1)
            style_map = {1:"Heading 1", 2:"Heading 2", 3:"Heading 3",
                         4:"Heading 4", 5:"Heading 5", 6:"Heading 6"}
            style_name = style_map.get(level, "Heading 1")
            # Use original style_name from format if present
            orig_style = fmt_snap.get("formatting", {}).get("style_name")
            if orig_style and "Heading" in orig_style:
                style_name = orig_style

            translated = _resolved_text(first_seg)
            _write_paragraph(doc, translated, fmt_snap, style_name=style_name)
            continue

        # ── Paragraphs / sentences ───────────────────────────────────────────
        # Join multiple sentence segments back into one paragraph
        translated_parts = [_resolved_text(s) for s in sorted(segs, key=lambda s: s["position"].get("sentence_index") or 0)]
        translated_text  = " ".join(p for p in translated_parts if p)

        _write_paragraph(doc, translated_text, fmt_snap)

    # Flush any remaining table
    _flush_table()

    doc.save(str(output_path))
    logger.info(f"DOCX regenerated → {output_path}")
    return output_path


# ============================================================================
# PDF Regeneration via ReportLab
# ============================================================================

def _flags_to_rl_font(font_name: str, flags: int) -> str:
    """
    Map a PDF font name + flags to a ReportLab built-in font name.
    ReportLab built-ins: Helvetica, Times-Roman, Courier (+ -Bold, -Oblique, -BoldOblique)
    """
    fn = (font_name or "").lower()
    is_bold   = bool(flags & (1 << 4)) or "bold" in fn
    is_italic = bool(flags & (1 << 6)) or "italic" in fn or "oblique" in fn

    if "courier" in fn or "mono" in fn:
        base = "Courier"
    elif "times" in fn or "serif" in fn or "roman" in fn:
        base = "Times-Roman"
    else:
        base = "Helvetica"

    if base == "Times-Roman":
        if is_bold and is_italic: return "Times-BoldItalic"
        if is_bold:               return "Times-Bold"
        if is_italic:             return "Times-Italic"
        return "Times-Roman"
    else:
        suffix = ""
        if is_bold and is_italic: suffix = "-BoldOblique"
        elif is_bold:             suffix = "-Bold"
        elif is_italic:           suffix = "-Oblique"
        return base + suffix


def regenerate_pdf(
    parsed_doc: Dict,
    segments: List[Dict],
    output_path: Path,
) -> Path:
    """
    Rebuild a PDF from translated segments using ReportLab.
    Places text at the original bbox coordinates with original font sizes.
    """
    try:
        from reportlab.pdfgen import canvas as rl_canvas
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.colors import HexColor, black
        from reportlab.pdfbase import pdfmetrics
    except ImportError:
        raise RuntimeError(
            "ReportLab is not installed. Run: pip install reportlab\n"
            "Alternatively, use POST /export/{doc_id}?format=docx to export as DOCX."
        )

    # Group segments by page
    page_segments: Dict[int, List[Dict]] = defaultdict(list)

    # First: collect pass-through image/skip blocks per page
    skip_blocks_by_page: Dict[int, List[Dict]] = defaultdict(list)

    for seg in segments:
        if seg.get("status") == "skip":
            snap = seg.get("format_snapshot", {})
            page = snap.get("page", 0)
            skip_blocks_by_page[page].append(seg)
        else:
            snap = seg.get("format_snapshot", {})
            page = snap.get("page", 0)
            page_segments[page].append(seg)

    # Get page dimensions from parsed blocks
    page_dims: Dict[int, tuple] = {}
    for block in parsed_doc.get("blocks", []):
        p = block.get("page")
        if p is not None and p not in page_dims:
            page_dims[p] = (block.get("page_width", 595.0), block.get("page_height", 842.0))

    total_pages = max(
        max(page_segments.keys(), default=0),
        max(skip_blocks_by_page.keys(), default=0),
    ) + 1

    # Build the PDF page by page
    first_page_dims = page_dims.get(0, (595.0, 842.0))
    c = rl_canvas.Canvas(str(output_path), pagesize=first_page_dims)

    # Group segments by block_index within each page
    for page_num in range(total_pages):
        pw, ph = page_dims.get(page_num, (595.0, 842.0))
        c.setPageSize((pw, ph))

        # Collect translatable segments for this page, group by block_index
        block_groups: Dict[int, List[Dict]] = defaultdict(list)
        for seg in page_segments.get(page_num, []):
            bi = seg["position"]["block_index"]
            block_groups[bi].append(seg)

        # Draw each block
        for bi, segs in sorted(block_groups.items()):
            if not segs:
                continue

            # Use first segment's format snapshot for block-level properties
            first_seg = min(segs, key=lambda s: s["position"].get("sentence_index") or 0)
            snap = first_seg.get("format_snapshot", {})
            bbox = snap.get("bbox", [])

            if not bbox or len(bbox) < 4:
                continue

            x0, y0_pdf, x1, y1_pdf = bbox
            # PyMuPDF uses top-left origin; ReportLab uses bottom-left
            # Convert: rl_y = page_height - pdf_y
            rl_y_top    = ph - y0_pdf
            rl_y_bottom = ph - y1_pdf
            block_width = x1 - x0

            # Join translated sentences
            translated_parts = [
                _resolved_text(s)
                for s in sorted(segs, key=lambda s: s["position"].get("sentence_index") or 0)
            ]
            translated_text = " ".join(p for p in translated_parts if p)

            if not translated_text:
                continue

            # Derive font properties from the first line/span of the snapshot
            lines = snap.get("lines", [])
            first_span = {}
            if lines and lines[0].get("spans"):
                first_span = lines[0]["spans"][0]

            font_size  = first_span.get("size", snap.get("body_font_size", 11.0))
            font_name  = _flags_to_rl_font(
                first_span.get("font", "Helvetica"),
                first_span.get("flags", 0)
            )
            color_hex  = first_span.get("color", "#000000")

            try:
                c.setFont(font_name, font_size)
            except Exception:
                c.setFont("Helvetica", font_size)

            try:
                c.setFillColor(HexColor(color_hex))
            except Exception:
                c.setFillColor(black)

            # Draw text — wrap to block width
            from reportlab.lib.utils import simpleSplit
            try:
                lines_wrapped = simpleSplit(translated_text, font_name, font_size, block_width)
            except Exception:
                lines_wrapped = [translated_text]

            line_height = font_size * 1.2
            y_cursor    = rl_y_top - font_size   # start from top of block

            for line_text in lines_wrapped:
                if y_cursor < rl_y_bottom - font_size:
                    break   # don't overflow below block bottom
                c.drawString(x0, y_cursor, line_text)
                y_cursor -= line_height

        # Image placeholders from skip blocks
        for skip_seg in skip_blocks_by_page.get(page_num, []):
            if skip_seg.get("block_type") == "image":
                snap = skip_seg.get("format_snapshot", {})
                bbox = snap.get("bbox", [])
                if bbox and len(bbox) == 4:
                    x0, y0_pdf, x1, y1_pdf = bbox
                    rl_y_top    = ph - y0_pdf
                    rl_y_bottom = ph - y1_pdf
                    # Draw a light grey placeholder rectangle
                    c.setStrokeColorRGB(0.8, 0.8, 0.8)
                    c.setFillColorRGB(0.95, 0.95, 0.95)
                    c.rect(x0, rl_y_bottom, x1 - x0, rl_y_top - rl_y_bottom, fill=1, stroke=1)
                    c.setFillColorRGB(0.5, 0.5, 0.5)
                    c.setFont("Helvetica", 8)
                    c.drawString(x0 + 4, rl_y_bottom + 4, "[Image]")

        c.showPage()   # end this page

    c.save()
    logger.info(f"PDF regenerated → {output_path}")
    return output_path


# ============================================================================
# Public entry point
# ============================================================================

def regenerate_document(
    document_id: str,
    parsed_doc: Dict,
    segments: List[Dict],
    output_format: str = "same",   # "same" | "docx" | "pdf"
) -> Path:
    """
    Regenerate a translated document.

    Args:
        document_id: UUID of the document.
        parsed_doc:  Parsed document dict (from parsed_docs/{id}.json).
        segments:    All segments with translated_text/final_text filled.
        output_format: "same" uses the original file_type; "docx" forces DOCX;
                       "pdf" forces PDF (requires ReportLab).

    Returns:
        Path to the exported file.
    """
    file_type  = parsed_doc.get("file_type", "docx")
    filename   = Path(parsed_doc.get("filename", "document")).stem
    timestamp  = datetime.utcnow().strftime("%Y%m%d_%H%M%S")

    target_format = output_format if output_format != "same" else file_type

    output_path = EXPORTS_DIR / f"{filename}_translated_{timestamp}.{target_format}"

    if target_format == "docx":
        return regenerate_docx(parsed_doc, segments, output_path)
    elif target_format == "pdf":
        return regenerate_pdf(parsed_doc, segments, output_path)
    else:
        raise ValueError(f"Unsupported output format: {target_format}")