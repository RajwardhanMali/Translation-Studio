"""
Document parser service — v2 (Format-Preserving).

Core design principle:
  Every block stores not just its text, but enough metadata to RECONSTRUCT
  the original document layout faithfully after translation.

DOCX blocks store:
  - style_name, alignment, indent levels, numbering info
  - per-run font properties (bold, italic, underline, size, color, font family)
  - table structure: row/col counts, merged cells, column widths
  - list level and numId for bullet/numbered lists

PDF blocks store:
  - exact bounding box (x0,y0,x1,y1) on each page
  - font name, size, color, flags (bold/italic) per span
  - page number, page dimensions (width/height)
  - block reading order (y then x sort)
  - line break positions within a block
  - image blocks: bbox, page number (for layout reconstruction)

This metadata is saved in parsed_docs/{id}.json and used by the
document regeneration service to rebuild the file after translation.
"""

import logging
import uuid
import re
from pathlib import Path
from typing import List, Dict, Any, Optional
from app.utils.file_handler import get_image_path
from app.services.ocr_engine import extract_text_from_image

logger = logging.getLogger(__name__)


# ============================================================================
# DOCX PARSER
# ============================================================================

def _rgb_to_hex(rgb_tuple) -> Optional[str]:
    """Convert (r,g,b) tuple to '#rrggbb' hex string."""
    if rgb_tuple is None:
        return None
    try:
        r, g, b = int(rgb_tuple[0]), int(rgb_tuple[1]), int(rgb_tuple[2])
        return f"#{r:02x}{g:02x}{b:02x}"
    except Exception:
        return None


def _pt_to_float(pt_val) -> Optional[float]:
    """Convert docx Pt/Emu length to float points, or None."""
    if pt_val is None:
        return None
    try:
        return float(pt_val.pt)
    except Exception:
        try:
            return float(pt_val)
        except Exception:
            return None


def _extract_run_formatting(run) -> Dict[str, Any]:
    """
    Extract all formatting properties from a docx Run object.
    Returns a dict that can reproduce the run's appearance.
    """
    fmt: Dict[str, Any] = {}

    try:
        fmt["bold"]      = run.bold
        fmt["italic"]    = run.italic
        fmt["underline"] = run.underline

        if run.font:
            fmt["font_name"]  = run.font.name
            fmt["font_size"]  = _pt_to_float(run.font.size)
            fmt["strike"]     = run.font.strike
            fmt["superscript"] = run.font.superscript
            fmt["subscript"]  = run.font.subscript
            fmt["all_caps"]   = run.font.all_caps
            fmt["small_caps"] = run.font.small_caps
            fmt["highlight_color"] = str(run.font.highlight_color) if run.font.highlight_color else None

            # Color
            if run.font.color and run.font.color.type is not None:
                try:
                    rgb = run.font.color.rgb
                    fmt["color"] = str(rgb)   # '000000' hex string from docx
                except Exception:
                    fmt["color"] = None
            else:
                fmt["color"] = None
    except Exception as e:
        logger.debug(f"Run formatting extraction partial error: {e}")

    return {k: v for k, v in fmt.items() if v is not None}


def _extract_paragraph_formatting(para) -> Dict[str, Any]:
    """
    Extract paragraph-level formatting: alignment, spacing, indent, style.
    """
    fmt: Dict[str, Any] = {}

    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        pf = para.paragraph_format

        fmt["style_name"]     = para.style.name if para.style else "Normal"
        fmt["alignment"]      = str(para.alignment) if para.alignment else None
        fmt["left_indent"]    = _pt_to_float(pf.left_indent)   if pf else None
        fmt["right_indent"]   = _pt_to_float(pf.right_indent)  if pf else None
        fmt["first_line_indent"] = _pt_to_float(pf.first_line_indent) if pf else None
        fmt["space_before"]   = _pt_to_float(pf.space_before)  if pf else None
        fmt["space_after"]    = _pt_to_float(pf.space_after)   if pf else None
        fmt["line_spacing"]   = _pt_to_float(pf.line_spacing)  if pf else None
        fmt["keep_together"]  = pf.keep_together if pf else None
        fmt["keep_with_next"] = pf.keep_with_next if pf else None
        fmt["page_break_before"] = pf.page_break_before if pf else None

        # Numbering / list info
        pPr = para._p.pPr if para._p is not None else None
        if pPr is not None:
            numPr = pPr.find(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr"
            )
            if numPr is not None:
                ilvl = numPr.find(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl"
                )
                numId = numPr.find(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId"
                )
                ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                fmt["list_level"] = int(ilvl.get(f"{{{ns}}}val", 0)) if ilvl is not None else 0
                fmt["list_num_id"] = int(numId.get(f"{{{ns}}}val", 0)) if numId is not None else None

    except Exception as e:
        logger.debug(f"Paragraph formatting partial error: {e}")

    return {k: v for k, v in fmt.items() if v is not None}


def parse_docx(filepath: Path, document_id: str) -> List[Dict[str, Any]]:
    """
    Parse a DOCX file into richly-formatted blocks.
    Each block includes enough metadata to reconstruct the document after translation.
    """
    try:
        from docx import Document
        from docx.table import Table as DocxTable
        from docx.text.paragraph import Paragraph as DocxParagraph
    except ImportError:
        raise RuntimeError("python-docx is not installed.")

    doc = Document(str(filepath))
    blocks: List[Dict] = []
    block_index = 0
    table_index = 0

    body = doc.element.body

    for child in body.iterchildren():
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        # ── Paragraphs / Headings ────────────────────────────────────────────
        if tag == "p":
            para = DocxParagraph(child, doc)
            
            # Extract images inline
            images_in_para = []
            if getattr(doc.part, "related_parts", None):
                for run in para.runs:
                    xml = run._element.xml
                    if "w:drawing" in xml:
                        match = re.search(r'r:embed="([^"]+)"', xml)
                        if match:
                            rid = match.group(1)
                            part = doc.part.related_parts.get(rid)
                            if part and part.content_type.startswith("image/"):
                                ext = part.content_type.split("/")[-1]
                                img_name = f"img_{document_id}_{block_index}_{len(images_in_para)}.{ext}"
                                img_path = get_image_path(document_id, img_name)
                                with open(img_path, "wb") as f:
                                    f.write(part.blob)
                                images_in_para.append(img_path)

            for img_path in images_in_para:
                ocr_text = extract_text_from_image(img_path)
                blocks.append({
                    "id": str(uuid.uuid4()),
                    "document_id": document_id,
                    "block_type": "image",
                    "text": ocr_text,
                    "src": str(img_path),
                    "position": {"block_index": block_index, "sentence_index": None, "phrase_index": None},
                })
                block_index += 1

            # Collect runs with per-run formatting
            runs_data = []
            full_text_parts = []
            for run in para.runs:
                run_text = run.text
                if run_text:
                    full_text_parts.append(run_text)
                    runs_data.append({
                        "text":       run_text,
                        "formatting": _extract_run_formatting(run),
                    })

            full_text = "".join(full_text_parts).strip()
            if not full_text:
                # Preserve empty paragraphs as spacers for layout
                blocks.append({
                    "id":           str(uuid.uuid4()),
                    "document_id":  document_id,
                    "block_type":   "spacer",
                    "text":         "",
                    "runs":         [],
                    "formatting":   _extract_paragraph_formatting(para),
                    "position":     {"block_index": block_index, "sentence_index": None, "phrase_index": None},
                })
                block_index += 1
                continue

            style_name = (para.style.name or "").lower() if para.style else ""
            para_fmt   = _extract_paragraph_formatting(para)

            if "heading" in style_name:
                try:
                    level = int(style_name.replace("heading", "").strip())
                except ValueError:
                    level = 1
                blocks.append({
                    "id":          str(uuid.uuid4()),
                    "document_id": document_id,
                    "block_type":  "heading",
                    "text":        full_text,
                    "runs":        runs_data,
                    "formatting":  para_fmt,
                    "level":       level,
                    "position":    {"block_index": block_index, "sentence_index": None, "phrase_index": None},
                })
            else:
                blocks.append({
                    "id":          str(uuid.uuid4()),
                    "document_id": document_id,
                    "block_type":  "paragraph",
                    "text":        full_text,
                    "runs":        runs_data,
                    "formatting":  para_fmt,
                    "position":    {"block_index": block_index, "sentence_index": None, "phrase_index": None},
                })
            block_index += 1

        # ── Tables ──────────────────────────────────────────────────────────
        elif tag == "tbl":
            tbl = DocxTable(child, doc)

            # Capture column widths
            col_widths = []
            try:
                for col in tbl.columns:
                    col_widths.append(_pt_to_float(col.width))
            except Exception:
                col_widths = []

            # Emit a table-start marker block
            table_block_id = str(uuid.uuid4())
            blocks.append({
                "id":           table_block_id,
                "document_id":  document_id,
                "block_type":   "table_start",
                "text":         "",
                "table_index":  table_index,
                "row_count":    len(tbl.rows),
                "col_count":    len(tbl.columns),
                "col_widths":   col_widths,
                "position":     {"block_index": block_index, "sentence_index": None, "phrase_index": None},
            })
            block_index += 1

            seen_tc = set()
            for r_idx, row in enumerate(tbl.rows):
                for c_idx, cell in enumerate(row.cells):
                    tc = cell._tc
                    if tc in seen_tc:
                        continue
                    seen_tc.add(tc)
                    
                    # Collect paragraphs inside each cell
                    cell_runs = []
                    cell_text_parts = []
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.text:
                                cell_text_parts.append(run.text)
                                cell_runs.append({
                                    "text":       run.text,
                                    "formatting": _extract_run_formatting(run),
                                })
                        cell_text_parts.append("\n")   # paragraph break within cell

                    cell_text = "".join(cell_text_parts).strip()

                    blocks.append({
                        "id":           str(uuid.uuid4()),
                        "document_id":  document_id,
                        "block_type":   "table_cell",
                        "text":         cell_text,
                        "runs":         cell_runs,
                        "table_index":  table_index,
                        "table_block_id": table_block_id,
                        "row":          r_idx,
                        "col":          c_idx,
                        "formatting":   {},
                        "position":     {"block_index": block_index, "sentence_index": None, "phrase_index": None},
                    })
                    block_index += 1

            # Table-end marker
            blocks.append({
                "id":          str(uuid.uuid4()),
                "document_id": document_id,
                "block_type":  "table_end",
                "text":        "",
                "table_index": table_index,
                "position":    {"block_index": block_index, "sentence_index": None, "phrase_index": None},
            })
            block_index += 1
            table_index += 1

    logger.info(f"DOCX parsed: {len(blocks)} blocks (format-preserving) from {filepath.name}")
    return blocks


# ============================================================================
# PDF PARSER — Format-Preserving
# ============================================================================

def _classify_block_role(
    spans: List[Dict],
    page_body_font_size: float,
    page_width: float,
    block_bbox: tuple,
) -> str:
    """
    Heuristically classify a PDF text block as heading/paragraph/caption/footer.
    Uses font size relative to body text, position on page, and text length.
    """
    if not spans:
        return "paragraph"

    sizes   = [s["size"] for s in spans if s.get("size", 0) > 0]
    avg_size = sum(sizes) / len(sizes) if sizes else page_body_font_size
    flags   = [s.get("flags", 0) for s in spans]
    is_bold  = any(f & 2**4 for f in flags)   # PyMuPDF bold flag is bit 4

    text = " ".join(s.get("text", "") for s in spans).strip()
    text_len = len(text)

    x0, y0, x1, y1 = block_bbox

    # Footer / header heuristic: very top or very bottom of page, short text
    if y1 < 60 or y0 > 740:   # approximate for A4 / Letter
        if text_len < 80:
            return "header_footer"

    # Heading: significantly larger than body, or bold + short
    if avg_size >= page_body_font_size + 3:
        return "heading"
    if is_bold and text_len < 120:
        return "heading"

    # Caption: small text, short, often under/above figures
    if avg_size < page_body_font_size - 1 and text_len < 200:
        return "caption"

    return "paragraph"


def _detect_body_font_size(page_blocks: List[Dict]) -> float:
    """
    Detect the most common font size on a page — this is the body text size.
    Used as baseline for heading detection.
    """
    from collections import Counter
    sizes = []
    for blk in page_blocks:
        if blk.get("type") != 0:
            continue
        for line in blk.get("lines", []):
            for span in line.get("spans", []):
                sz = round(span.get("size", 0), 1)
                if sz > 0:
                    sizes.append(sz)
    if not sizes:
        return 11.0
    most_common = Counter(sizes).most_common(1)
    return most_common[0][0] if most_common else 11.0


def parse_pdf(filepath: Path, document_id: str) -> List[Dict[str, Any]]:
    """
    Parse a PDF with full format preservation.

    Each block stores:
      - page_number, page_width, page_height
      - bbox (x0, y0, x1, y1) — exact position for reconstruction
      - lines: list of {spans: [{text, font, size, color, flags, bbox}]}
      - block_role: heading | paragraph | caption | header_footer | image
      - reading_order index (y-primary, x-secondary sort within page)
    """
    try:
        import fitz
    except ImportError:
        raise RuntimeError("PyMuPDF (fitz) is not installed.")

    doc   = fitz.open(str(filepath))
    blocks: List[Dict] = []
    block_index = 0

    for page_num, page in enumerate(doc):
        page_rect   = page.rect
        page_width  = page_rect.width
        page_height = page_rect.height

        text_dict   = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)
        page_blocks = text_dict.get("blocks", [])

        # Detect body font size for this page
        body_size   = _detect_body_font_size(page_blocks)

        # Sort blocks in reading order: top→bottom, left→right
        text_blocks = [b for b in page_blocks if b.get("type") == 0]
        image_blocks= [b for b in page_blocks if b.get("type") == 1]

        # ── Table blocks ────────────────────────────────────────────────────
        tabs = page.find_tables()
        table_bboxes = []
        if tabs:
            for t_idx, tab in enumerate(tabs.tables):
                table_bboxes.append(tab.bbox)
                
                table_block_id = str(uuid.uuid4())
                blocks.append({
                    "id":           table_block_id,
                    "document_id":  document_id,
                    "block_type":   "table_start",
                    "text":         "",
                    "table_index":  t_idx,
                    "row_count":    len(tab.cells) if tab.cells else 0,
                    "col_count":    len(tab.cells[0]) if tab.cells and tab.cells[0] else 0,
                    "page":         page_num,
                    "position":     {"block_index": block_index, "sentence_index": None, "phrase_index": None},
                })
                block_index += 1
                
                for r_idx, row in enumerate(tab.extract() or []):
                    for c_idx, cell_text in enumerate(row or []):
                        if cell_text is None:
                            cell_text = ""
                        blocks.append({
                            "id":           str(uuid.uuid4()),
                            "document_id":  document_id,
                            "block_type":   "table_cell",
                            "text":         cell_text.strip().replace("\n", " "),
                            "table_index":  t_idx,
                            "table_block_id": table_block_id,
                            "row":          r_idx,
                            "col":          c_idx,
                            "formatting":   {},
                            "position":     {"block_index": block_index, "sentence_index": None, "phrase_index": None},
                        })
                        block_index += 1
                        
                blocks.append({
                    "id":          str(uuid.uuid4()),
                    "document_id": document_id,
                    "block_type":  "table_end",
                    "text":        "",
                    "table_index": t_idx,
                    "position":    {"block_index": block_index, "sentence_index": None, "phrase_index": None},
                })
                block_index += 1

        # Sort text blocks in reading order: top→bottom, left→right
        text_blocks.sort(key=lambda b: (round(b["bbox"][1] / 20) * 20, b["bbox"][0]))

        # ── Image blocks ────────────────────────────────────────────────────
        for img_idx, img_blk in enumerate(image_blocks):
            bbox = img_blk.get("bbox", [0, 0, 0, 0])
            img_bytes = img_blk.get("image")
            img_path_str = ""
            ocr_text = ""
            
            if img_bytes:
                img_ext = img_blk.get("ext", "png")
                img_name = f"img_{document_id}_{page_num}_{img_idx}.{img_ext}"
                try:
                    img_path = get_image_path(document_id, img_name)
                    with open(img_path, "wb") as f:
                        f.write(img_bytes)
                    img_path_str = str(img_path)
                    ocr_text = extract_text_from_image(img_path)
                except Exception as e:
                    logger.warning(f"Failed to process image: {e}")

            blocks.append({
                "id":           str(uuid.uuid4()),
                "document_id":  document_id,
                "block_type":   "image",
                "text":         ocr_text,
                "src":          img_path_str,
                "page":         page_num,
                "page_width":   page_width,
                "page_height":  page_height,
                "bbox":         list(bbox),
                "position":     {"block_index": block_index, "sentence_index": None, "phrase_index": None},
            })
            block_index += 1

        # Helper to check if block is in table
        def _in_table(bx):
            bx0, by0, bx1, by1 = bx
            cx, cy = (bx0+bx1)/2, (by0+by1)/2
            for tx0, ty0, tx1, ty1 in table_bboxes:
                if tx0 <= cx <= tx1 and ty0 <= cy <= ty1:
                    return True
            return False

        # ── Text blocks ─────────────────────────────────────────────────────
        for pdf_block in text_blocks:
            if _in_table(pdf_block.get("bbox", [0, 0, 0, 0])):
                continue
            
            lines_data = []
            all_spans  = []
            full_text_parts = []

            for line in pdf_block.get("lines", []):
                line_spans = []
                for span in line.get("spans", []):
                    span_text = span.get("text", "")
                    if not span_text:
                        continue

                    # Extract color as hex
                    raw_color = span.get("color", 0)
                    if isinstance(raw_color, int):
                        r = (raw_color >> 16) & 0xFF
                        g = (raw_color >>  8) & 0xFF
                        b =  raw_color        & 0xFF
                        color_hex = f"#{r:02x}{g:02x}{b:02x}"
                    else:
                        color_hex = "#000000"

                    span_data = {
                        "text":   span_text,
                        "font":   span.get("font", ""),
                        "size":   round(span.get("size", 11.0), 2),
                        "color":  color_hex,
                        "flags":  span.get("flags", 0),
                        "bbox":   list(span.get("bbox", [0, 0, 0, 0])),
                        "origin": list(span.get("origin", [0, 0])),
                    }
                    line_spans.append(span_data)
                    all_spans.append(span_data)
                    full_text_parts.append(span_text)

                if line_spans:
                    lines_data.append({
                        "spans": line_spans,
                        "bbox":  list(line.get("bbox", [0, 0, 0, 0])),
                        "wmode": line.get("wmode", 0),
                        "dir":   list(line.get("dir", [1, 0])),
                    })

            full_text = " ".join(full_text_parts).strip()
            # Collapse excessive whitespace from PDF extraction artifacts
            full_text = re.sub(r"  +", " ", full_text)

            if not full_text:
                continue

            bbox = pdf_block.get("bbox", [0, 0, page_width, 12])
            block_role = _classify_block_role(all_spans, body_size, page_width, tuple(bbox))

            # Determine heading level from font size
            level = None
            if block_role == "heading":
                avg_size = sum(s["size"] for s in all_spans) / len(all_spans) if all_spans else body_size
                if avg_size >= body_size + 10:
                    level = 1
                elif avg_size >= body_size + 6:
                    level = 2
                elif avg_size >= body_size + 3:
                    level = 3
                else:
                    level = 4

            block = {
                "id":           str(uuid.uuid4()),
                "document_id":  document_id,
                "block_type":   block_role,
                "text":         full_text,
                "lines":        lines_data,
                "page":         page_num,
                "page_width":   round(page_width, 2),
                "page_height":  round(page_height, 2),
                "bbox":         [round(v, 2) for v in bbox],
                "body_font_size": body_size,
                "position":     {"block_index": block_index, "sentence_index": None, "phrase_index": None},
            }
            if level is not None:
                block["level"] = level

            blocks.append(block)
            block_index += 1

    doc.close()
    logger.info(f"PDF parsed (format-preserving): {len(blocks)} blocks from {filepath.name}")
    return blocks


# ============================================================================
# Public entry point
# ============================================================================

def parse_document(filepath: Path, document_id: str, file_type: str) -> List[Dict[str, Any]]:
    if file_type == "docx":
        return parse_docx(filepath, document_id)
    elif file_type == "pdf":
        return parse_pdf(filepath, document_id)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")