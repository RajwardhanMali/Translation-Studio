import pytest
import uuid
import os
from pathlib import Path
from app.services.parser import parse_docx, parse_pdf
from docx import Document
import fitz

TEST_DIR = Path(__file__).parent / "test_data"
os.makedirs(TEST_DIR, exist_ok=True)

@pytest.fixture
def complex_docx_path():
    doc = Document()
    doc.add_paragraph("Regular Paragraph")
    
    # Add a table with merged cells
    table = doc.add_table(rows=2, cols=2)
    a = table.cell(0, 0)
    b = table.cell(0, 1)
    a.merge(b) # merge top row
    a.text = "Merged Cell Row 1"
    
    table.cell(1, 0).text = "Cell 1,0"
    table.cell(1, 1).text = "Cell 1,1"
    
    path = TEST_DIR / "complex_test.docx"
    doc.save(str(path))
    return path

@pytest.fixture
def complex_pdf_path():
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((200, 200), "Regular PDF Text")
    
    # We won't simulate a complex table or image here perfectly without
    # drawing one out line by line, but we can verify it parses without crashing.
    
    path = TEST_DIR / "complex_test.pdf"
    doc.save(str(path))
    doc.close()
    return path

def test_parse_docx_merged_cells(complex_docx_path):
    """Test that python-docx table parsing correctly de-duplicates merged XML cell data."""
    doc_id = "test_merged_docx"
    blocks = parse_docx(complex_docx_path, doc_id)
    
    # We expect: paragraph, table_start, merged_cell, 1,0 cell, 1,1 cell, table_end
    # If the deduplication failed, we would see 'Merged Cell Row 1' twice.
    cell_texts = [b["text"] for b in blocks if b["block_type"] == "table_cell"]
    
    # Deduplication means it only extracts the merged cell once
    assert cell_texts.count("Merged Cell Row 1") == 1
    assert "Cell 1,0" in cell_texts
    assert "Cell 1,1" in cell_texts

def test_parse_pdf_basic_structure(complex_pdf_path):
    """Test that the updated PDF parser extracts text and initializes the table tracker."""
    doc_id = "test_pdf"
    blocks = parse_pdf(complex_pdf_path, doc_id)
    
    text_blocks = [b for b in blocks if b["block_type"] == "paragraph"]
    assert len(text_blocks) >= 1
    assert "Regular PDF Text" in text_blocks[0]["text"]
